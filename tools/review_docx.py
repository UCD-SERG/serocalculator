from __future__ import annotations

import copy
import datetime as dt
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree


AUTHOR = "Codex editorial review"
DATE = dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat().replace(
    "+00:00", "Z"
)
SOURCE_DIR = Path(
    "/Users/ezramorrison/Library/CloudStorage/OneDrive-UCDavisHealth/Work/"
    "Projects/SeroEpi/Serocalculator/Publications/revision 1"
)
OUTPUT_DIR = Path(
    "/Users/ezramorrison/Documents/GitHub/serocalculator/reviewed-documents"
)


def paragraph_text_with_math(paragraph) -> str:
    texts = paragraph._p.xpath(".//w:t | .//m:t")
    return "".join(node.text or "" for node in texts)


def next_revision_id(document: Document) -> int:
    ids = []
    for node in document.element.xpath(".//w:ins | .//w:del"):
        value = node.get(qn("w:id"))
        if value is not None:
            try:
                ids.append(int(value))
            except ValueError:
                pass
    return max(ids, default=0) + 1


def revision_element(tag: str, revision_id: int) -> OxmlElement:
    element = OxmlElement(tag)
    element.set(qn("w:id"), str(revision_id))
    element.set(qn("w:author"), AUTHOR)
    element.set(qn("w:date"), DATE)
    return element


def text_run(text: str, deleted: bool = False, bold: bool = False) -> OxmlElement:
    run = OxmlElement("w:r")
    if bold:
        properties = OxmlElement("w:rPr")
        properties.append(OxmlElement("w:b"))
        run.append(properties)
    text_element = OxmlElement("w:delText" if deleted else "w:t")
    if text.startswith(" ") or text.endswith(" ") or "  " in text:
        text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    run.append(text_element)
    return run


def replace_paragraph(paragraph, new_text: str, revision_id: int) -> int:
    old_text = paragraph_text_with_math(paragraph)
    if not old_text.strip():
        raise ValueError("Refusing to replace a blank paragraph")

    paragraph_xml = paragraph._p
    properties = paragraph_xml.find(qn("w:pPr"))
    for child in list(paragraph_xml):
        if child is not properties:
            paragraph_xml.remove(child)

    deletion = revision_element("w:del", revision_id)
    deletion.append(text_run(old_text, deleted=True))
    paragraph_xml.append(deletion)
    revision_id += 1

    insertion = revision_element("w:ins", revision_id)
    if new_text.startswith("Response:"):
        insertion.append(text_run("Response:", bold=True))
        insertion.append(text_run(new_text[len("Response:") :]))
    else:
        insertion.append(text_run(new_text))
    paragraph_xml.append(insertion)
    return revision_id + 1


def replace_text_fragment(paragraph, old_text: str, new_text: str,
                          revision_id: int) -> int:
    candidates = [
        node for node in paragraph._p.xpath(".//w:t")
        if node.text is not None and old_text in node.text
    ]
    if len(candidates) != 1:
        raise ValueError(
            f"Expected one text node containing {old_text!r}; found {len(candidates)}"
        )
    text_node = candidates[0]
    run = text_node.getparent()
    if run.tag != qn("w:r") or len(run.xpath("./w:t")) != 1:
        raise ValueError(f"Unsupported run structure for {old_text!r}")
    parent = run.getparent()
    position = parent.index(run)
    before, after = text_node.text.split(old_text, 1)
    parent.remove(run)

    elements = []
    if before:
        elements.append(text_run(before))
    deletion = revision_element("w:del", revision_id)
    deletion.append(text_run(old_text, deleted=True))
    elements.append(deletion)
    revision_id += 1
    insertion = revision_element("w:ins", revision_id)
    insertion.append(text_run(new_text))
    elements.append(insertion)
    revision_id += 1
    if after:
        elements.append(text_run(after))
    for offset, element in enumerate(elements):
        parent.insert(position + offset, element)
    return revision_id


def replace_table_cell(document: Document, table_index: int, row: int, column: int,
                       new_text: str, revision_id: int) -> int:
    cell = document.tables[table_index].cell(row, column)
    nonblank = [p for p in cell.paragraphs if paragraph_text_with_math(p).strip()]
    paragraph = nonblank[0] if nonblank else cell.paragraphs[0]
    if not paragraph_text_with_math(paragraph).strip():
        run = paragraph.add_run("[blank]")
        run._r.getparent().remove(run._r)
        paragraph._p.append(text_run("[blank]"))
    return replace_paragraph(paragraph, new_text, revision_id)


def insert_table_cell(document: Document, table_index: int, row: int, column: int,
                      new_text: str, revision_id: int) -> int:
    cell = document.tables[table_index].cell(row, column)
    paragraph = cell.paragraphs[0]
    if paragraph_text_with_math(paragraph).strip():
        raise ValueError("Insertion-only table target is not blank")
    insertion = revision_element("w:ins", revision_id)
    insertion.append(text_run(new_text))
    paragraph._p.append(insertion)
    return revision_id + 1


def enable_track_revisions(document: Document) -> None:
    settings = document.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        settings.insert(0, OxmlElement("w:trackRevisions"))


def apply_paragraph_replacements(document: Document, replacements: dict[int, str],
                                 revision_id: int) -> int:
    for index, new_text in replacements.items():
        revision_id = replace_paragraph(document.paragraphs[index], new_text, revision_id)
    return revision_id


def save_and_validate(document: Document, output_path: Path) -> tuple[int, int]:
    enable_track_revisions(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    with zipfile.ZipFile(output_path) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise ValueError(f"Corrupt ZIP member: {bad_member}")
        document_xml = archive.read("word/document.xml")
        settings_xml = archive.read("word/settings.xml")

    root = etree.fromstring(document_xml)
    settings = etree.fromstring(settings_xml)
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    }
    insertions = len(root.xpath(".//w:ins", namespaces=namespaces))
    deletions = len(root.xpath(".//w:del", namespaces=namespaces))
    if not settings.xpath(".//w:trackRevisions", namespaces=namespaces):
        raise ValueError("Track revisions setting is missing")
    if insertions == 0 or deletions == 0:
        raise ValueError("Revision markup is missing")

    Document(output_path)
    return insertions, deletions


def review_manuscript() -> tuple[Path, int, int]:
    source = SOURCE_DIR / "Serocalculator manuscript_IJE_resubmit 2_ clean.docx"
    output = OUTPUT_DIR / "Serocalculator manuscript_IJE_resubmit 2_TRACKED.docx"
    document = Document(source)
    replacements = {
        28: (
            "Implementation: serocalculator is an open-source R package that uses a "
            "likelihood-based framework incorporating modeled antibody decay, "
            "between-person kinetic heterogeneity, cross-reactivity, and measurement "
            "noise to estimate seroincidence rates under a Poisson infection process "
            "from cross-sectional serological data."
        ),
        29: (
            "General features: The package supports overall and stratified "
            "seroincidence estimation using one or more biomarkers. It requires three "
            "inputs: (1) a pre-estimated seroresponse model characterizing "
            "post-infection antibody waning; (2) noise parameters describing "
            "cross-reactivity, non-specific binding, and assay error; and (3) "
            "quantitative antibody responses from a cross-sectional survey. It is "
            "computationally efficient, well documented, and includes a point-and-click "
            "R Shiny interface. These features promote use in research and public health."
        ),
        37: (
            "New methods estimate seroincidence rates from cross-sectional serosurveys "
            "by incorporating models of antibody decay dynamics from patients with "
            "confirmed infections (seroresponse parameters) 2,8,9. In brief, this "
            "approach uses maximum likelihood estimation (MLE) to determine the most "
            "likely seroincidence rate given the observed cross-sectional population "
            "data, modeled antibody decay trajectories, non-specific binding of assay "
            "probes, and laboratory assay variability 1,2,8–11. This method has "
            "previously been used to estimate the seroincidence of enteric fever, a "
            "systemic bacterial infection caused by Salmonella enterica serovars Typhi "
            "and Paratyphi. Although the resulting seroincidence estimates were "
            "substantially higher than blood-culture-based clinical incidence, they "
            "preserved the same rank order of clinical burden across sites 12."
        ),
        38: (
            "Here, we introduce serocalculator, a statistical software package that "
            "estimates seroincidence rates from cross-sectional serosurveys by "
            "integrating independent models of post-infection antibody kinetics with "
            "measurement-noise correction. A detailed list of requirements for the "
            "serocalculator approach is provided in Supplement S1 1,8–10. An earlier "
            "package, seroincidence 13, was developed and maintained from 2015 to 2018 "
            "but did not incorporate biological or measurement noise. serocalculator "
            "supersedes that package by incorporating these methodological advances, "
            "combining information across multiple antibody responses and isotypes, and "
            "adding functions for simulating cross-sectional data and modern "
            "ggplot2-based visualizations 14."
        ),
        45: (
            "Three inputs are required: (1) previously estimated longitudinal "
            "seroresponse parameters describing the antibody trajectories; (2) noise "
            "parameters describing cross-reactivity or non-specific binding and "
            "laboratory assay error; and (3) quantitative antibody levels from a "
            "cross-sectional population-based serosurvey. Users can access small "
            "example files within the package or import larger sample files from the "
            "publicly available Serocalculator Data Repository on Open Science "
            "Framework. All data must use the same assay and calibration. Further "
            "requirements, formatting guidance, and example inputs are provided in "
            "Supplements S1 and S3."
        ),
        53: (
            "The noise-parameters object represents two sources of variation distinct "
            "from between-person kinetic heterogeneity. Biological noise (nu) is "
            "additive signal from cross-reactivity or non-specific binding and is "
            "modeled as Uniform(0, nu); nu is commonly estimated as the 95th percentile "
            "of suitable negative controls. Measurement noise (eps) is multiplicative "
            "relative error modeled as Uniform(-eps, eps). Thus, eps is the maximum "
            "relative deviation, not a coefficient of variation (CV); under this "
            "uniform model, eps = sqrt(3) × CV. Lower and upper quantification limits "
            "may be specified as y.low and y.high, and observations outside those "
            "limits are treated as censored. Further details are in Supplement S6."
        ),
        55: (
            "Users can import noise parameters (CSV or RDS) as a data.frame and apply "
            "as_noise_params() to create a correctly formatted noise_params object. "
            "Files at a URL can instead be imported with load_noise_params(). Required "
            "variables are one or more antigen–antibody-isotype pairs matching the "
            "seroresponse parameters (antigen_iso), biological noise (nu), measurement "
            "noise (eps), and the lower and upper assay limits (y.low and y.high)."
        ),
        64: (
            "We demonstrate serocalculator with a reproducible enteric fever example "
            "using a subset of serosurvey data collected in Bangladesh, Nepal, and "
            "Pakistan during 2016–2021 12,17. Additional details are available in "
            "Supplement S7, and the data are publicly available in the Serocalculator "
            "Data Repository on Open Science Framework 18. This example is also "
            "presented as an article on the serocalculator website. Figure 2 provides "
            "the corresponding code and output."
        ),
        71: (
            "serocalculator provides a computationally efficient method for estimating "
            "population-level seroincidence rates from multiple biomarkers while "
            "accounting for antibody decay, between-person kinetic heterogeneity, "
            "cross-reactivity or non-specific binding, and assay noise. By updating the "
            "analytic methods in its predecessor 13 and integrating new tutorials, "
            "simulations, and visualizations, the package expands the accessibility of "
            "quantitative seroepidemiology. serocalculator has been applied to enteric "
            "fever and scrub typhus studies 12,19–21, and its methods are increasingly "
            "used by research teams and public health institutions to inform typhoid "
            "vaccine introduction and surveillance in endemic regions."
        ),
        73: (
            "In contrast, serocalculator is designed to estimate population-level "
            "seroincidence by pairing quantitative antibody responses with modeled "
            "antibody decay curves and marginalizing over between-person kinetic "
            "heterogeneity. It supports multiple antigen-isotype pairs and separately "
            "models cross-reactivity or non-specific binding and assay measurement "
            "error. Collectively, these packages expand access to seroepidemiologic "
            "tools, with serocalculator filling a distinct niche in cross-sectional "
            "serosurveillance. Future work will assess simplifying computational "
            "assumptions and structural modifications that relax them."
        ),
        89: "Supplementary materials are available online.",
    }
    revision_id = apply_paragraph_replacements(document, replacements,
                                                next_revision_id(document))
    revision_id = replace_text_fragment(
        document.paragraphs[48], "Supplement S4", "Supplement S5", revision_id
    )
    insertions, deletions = save_and_validate(document, output)
    return output, insertions, deletions


def review_supplement() -> tuple[Path, int, int]:
    source = SOURCE_DIR / "Supplementary material_serocalculator_resubmit_v2_clean.docx"
    output = OUTPUT_DIR / "Supplementary material_serocalculator_resubmit_v2_TRACKED.docx"
    document = Document(source)
    replacements = {
        23: (
            "The table shells below illustrate the three required serocalculator "
            "inputs. Any input may include stratifying variables, but only strata "
            "defined in the cross-sectional population data produce stratified "
            "seroincidence estimates. Noise and longitudinal seroresponse parameters "
            "may be specified by stratum or supplied as overall values applied to every "
            "cross-sectional stratum."
        ),
        24: "Table S3.1: Longitudinal seroresponse parameters",
        33: "Table S3.2: Noise parameters",
        38: (
            "eps = Bound on relative measurement error, range 0–1; eps is not the "
            "assay CV. Under the uniform-error model, eps = sqrt(3) × CV (see "
            "Supplement S6)."
        ),
        39: (
            "nu = Width of additive biological noise from cross-reactivity or "
            "non-specific binding; commonly estimated as the 95th percentile among "
            "suitable negative controls (see Supplement S6)."
        ),
        41: "Table S3.3: Cross-sectional population data",
        50: (
            "We do not currently recommend one minimum cross-sectional sample size, "
            "because the required sample depends on the research question, biomarker "
            "distributions, and heterogeneity in post-infection seroresponses. Users "
            "should use particular caution with stratified estimates based on small "
            "numbers (e.g., n < 20), which may be unstable. We recommend simulation to "
            "assess an existing study's precision or determine the sample size needed "
            "for a target effect size. The sim_pop_data() function supports such "
            "simulation-based assessments; see the Simulation Studies article on the "
            "package website."
        ),
        65: (
            "The full seroresponse is summarized by the five-parameter kinetic set. To "
            "account for between-person heterogeneity, serocalculator averages the "
            "likelihood over Monte Carlo draws from the kinetic-parameter distribution "
            "estimated from longitudinal confirmed-case data. This marginalization "
            "retains variation caused by random kinetic parameters as well as uncertainty "
            "represented by the supplied draws. In particular, variation in the waning "
            "rate can cause individual trajectories to diverge or converge over time, "
            "so the marginal population variance around the mean trajectory may depend "
            "on time since infection. This is distinct from residual observation noise: "
            "the upstream serodynamics model assumes constant residual variance on the "
            "log-antibody scale conditional on individual kinetic parameters, and "
            "serocalculator does not fit an additional residual-variance function of "
            "time. Instead, it carries trajectory-induced variation forward through "
            "Monte Carlo averaging. Its additive biological-noise width and relative "
            "measurement-noise width are also not explicit functions of time. 2,3,5"
        ),
        94: (
            "The longitudinal seroresponse parameters in Table S3.1 are components of "
            "the within-host model proposed by de Graaf et al. (2014), extended by "
            "Teunis et al. (2016) and Diekmann et al. (2018).1–3 The model has an "
            "active-infection/antibody-growth phase and a post-infection/antibody-decay "
            "phase. Beginning at infection, antibody concentrations grow exponentially "
            "until the peak response at time t1 (Figure S5.1a). This brief growth phase "
            "is ignored in the final seroincidence model. After t1, antibody concentrations "
            "follow the power-function decay phase shown in Figure S5.1b."
        ),
        107: (
            "This model is fitted in a Bayesian framework using JAGS, and posterior "
            "draws are then used in serocalculator. The companion R package "
            "serodynamics, available on CRAN, helps users model post-infection "
            "seroresponse parameters from longitudinal cohorts of confirmed cases. "
            "These models can be stratified by age, country, or other relevant "
            "characteristics when seroresponse dynamics differ among subpopulations."
        ),
        109: (
            "Figure S5.1a–b: Grey dots show longitudinal antibody measurements from "
            "confirmed cases, and grey lines connect measurements from the same person. "
            "The multicolored line shows the median modeled response; red indicates "
            "higher and blue lower quantitative responses. The blue dashed vertical "
            "line indicates time to peak (t1), at which the peak response is y1."
        ),
        143: (
            "serocalculator incorporates two sources of noise that are distinct from "
            "between-person kinetic heterogeneity: biological noise from cross-reactivity "
            "or non-specific binding and measurement noise from laboratory analysis. In "
            "the package, biological noise is denoted nu and measurement noise eps.1"
        ),
        144: (
            "Biological noise, nu, represents additive signal from cross-reactivity or "
            "non-specific binding. It is commonly estimated as the 95th percentile of "
            "responses to the antigen-isotype in a reference population with low or no "
            "recent exposure. The model approximates this signal with Uniform(0, nu). "
            "Teunis and van Eijkeren (2020) showed that incidence estimates are sensitive "
            "to the width of the noise distribution but comparatively robust to its "
            "shape; the particular use of the 95th percentile is an adopted convention, "
            "not an optimized cutoff. Biological noise is non-negative.1"
        ),
        146: (
            "Measurement noise, eps, represents multiplicative laboratory assay error. "
            "serocalculator models the relative error as Uniform(-eps, eps), so eps is "
            "the largest relative deviation rather than the coefficient of variation "
            "(CV). Under this model, CV = eps / sqrt(3), or equivalently eps = sqrt(3) × "
            "CV. The CV should ideally be estimated from replicates across plates rather "
            "than within one plate. Replicate measurements may reduce measurement error, "
            "but eps should correspond to the measurement process used for the analyzed "
            "cross-sectional values."
        ),
        147: (
            "Teunis and van Eijkeren (2020) provide the underlying noise model, and "
            "Aiemjoy et al. (2022) describe application-specific noise-estimation "
            "procedures.1,2"
        ),
        157: (
            "The SEES study demonstrated that, when combined with longitudinal antibody "
            "dynamics, cross-sectional responses to HlyE and LPS could estimate "
            "seroconversion rates and recover differences correlated with clinical "
            "incidence.5 Available data include IgA and IgG responses to S. Typhi and "
            "S. Paratyphi antigens HlyE and LPS from both confirmed enteric fever cases "
            "and population-based participants in the same catchment areas. Full details "
            "of the SEAP and SEES studies have been published elsewhere.5–7 The example "
            "does not include vaccination because typhoid conjugate vaccine targets Vi, "
            "a different antigen that does not cross-react with the HlyE and LPS assays."
        ),
    }
    revision_id = apply_paragraph_replacements(document, replacements,
                                                next_revision_id(document))
    revision_id = replace_text_fragment(
        document.paragraphs[64], "Supplement S1", "Supplement S5", revision_id
    )
    table_replacements = {
        (0, 5, 1): (
            "The model requires participant age to calculate the age-dependent "
            "probability of never having been infected, as introduced by Teunis and van "
            "Eijkeren (2020)."
        ),
        (0, 5, 2): (
            "The current implementation cannot calculate seroincidence when age is "
            "missing."
        ),
        (1, 1, 2): (
            "If the longitudinal cohort is not representative—for example, because of "
            "spectrum bias, mismatched ages, immunosuppression, or different exposure "
            "histories—seroincidence may be biased. If the cohort overrepresents severe "
            "cases with higher peaks, the model may overestimate time since infection "
            "and underestimate seroincidence (Simonsen et al., 2009; Aiemjoy et al., "
            "2022)."
        ),
        (1, 8, 1): (
            "The model uses biological noise to represent additive cross-reactivity or "
            "non-specific binding and measurement noise to represent multiplicative "
            "assay error. These parameters may vary by antigen-isotype and specified "
            "strata but are treated as fixed within each supplied noise-parameter row "
            "(Teunis & van Eijkeren, 2020)."
        ),
        (1, 10, 0): "10. Adequate Representation of Residual Cross-Reactivity",
        (1, 10, 1): (
            "After infection-specific antigen selection, remaining cross-reactive or "
            "non-specific signal is assumed to be adequately represented by the "
            "biological-noise parameter estimated from suitable negative controls "
            "(Aiemjoy et al., 2022)."
        ),
        (1, 10, 2): (
            "If residual cross-reactivity differs between the negative-control and "
            "target populations, biological noise may be misspecified. Underestimation "
            "can inflate seroincidence by attributing background signal to recent "
            "infection; overestimation can attenuate true infection signal."
        ),
    }
    for (table, row, column), new_text in table_replacements.items():
        revision_id = replace_table_cell(document, table, row, column, new_text,
                                         revision_id)
    table_insertions = {
        (0, 4, 0): "4. Availability of Noise-Parameter Data",
        (0, 4, 1): (
            "The model requires biological-noise and measurement-noise parameters for "
            "each antigen-isotype and applicable stratum. Biological noise should be "
            "estimated from suitable negative controls, and measurement noise from assay "
            "validation or replicate measurements."
        ),
        (0, 4, 2): (
            "Without suitable noise parameters, the model cannot distinguish background "
            "or assay variation from infection-related antibody signal. Misspecification "
            "can bias seroincidence, especially at low incidence and young ages."
        ),
    }
    for (table, row, column), new_text in table_insertions.items():
        revision_id = insert_table_cell(document, table, row, column, new_text,
                                        revision_id)
    insertions, deletions = save_and_validate(document, output)
    return output, insertions, deletions


def review_response_letter() -> tuple[Path, int, int]:
    source = SOURCE_DIR / "Serocalculator_Response to Reviewers.docx"
    output = OUTPUT_DIR / "Serocalculator_Response to Reviewers_TRACKED.docx"
    document = Document(source)
    replacements = {
        13: (
            "Even without the companion package, serocalculator had been downloaded "
            "from CRAN more than 36,000 times as of May 31, 2026, and we have taught "
            "multiple full-day workshops on use of these tools, indicating substantial "
            "interest in the international research community. We hope that publication "
            "and continued development of both packages will broaden their use across "
            "disease areas. Thank you again for your time and consideration. Our "
            "point-by-point responses follow."
        ),
        25: (
            "“The model assumes a constant, homogeneous force of infection (a Poisson "
            "process), no durable post-infection immunity, transportability of the "
            "longitudinal seroresponse distribution to the target population, and "
            "conditional independence among antigen-isotype responses. Supplement S2 "
            "describes these and other assumptions and their implications.”"
        ),
        26: "More details and implications have been added to the new Supplement S2.",
        29: (
            "Response: Thank you for highlighting this limitation. The current model "
            "assumes a relatively stable transmission environment. Strong seasonality, "
            "epidemic waves, or secular trends can make the estimate a time-averaged rate "
            "that does not represent current transmission. Sampling consistently over "
            "time, stratifying by period where justified, and simulation-based sensitivity "
            "analysis can help assess this limitation but do not remove the underlying "
            "assumption. We now state the assumptions in the Introduction and discuss "
            "their implications in Supplement S2:"
        ),
        30: (
            "“The model assumes a constant, homogeneous force of infection (a Poisson "
            "process), no durable post-infection immunity, transportability of the "
            "longitudinal seroresponse distribution to the target population, and "
            "conditional independence among antigen-isotype responses. Supplement S2 "
            "describes these and other assumptions and their implications.”"
        ),
        32: (
            "“Finally, the current method does not account for seasonality; relaxing the "
            "constant-incidence assumption is an area for future research.”"
        ),
        36: (
            "Response: Thank you for this helpful comment. We combined the previous "
            "figures into Figure 2 and created a new Figure 1 summarizing the "
            "serocalculator workflow. Although we initially omitted equations to match "
            "other IJE Software Application Profiles, we agree that a mathematical "
            "overview is useful. We therefore added key equations to Figure 1 and a full "
            "description of the mathematical framework in Supplement S4."
        ),
        56: (
            "Response: Thank you for pointing this out. We agree that conditional "
            "independence may be unrealistic for some diseases. Residual positive "
            "correlation among antigen-isotype responses means that multiplying their "
            "likelihoods can overstate the amount of independent information and produce "
            "standard errors that are too small; it is therefore not necessarily a "
            "conservative assumption. We are developing methods that model cross-biomarker "
            "correlation directly. The manuscript now identifies conditional independence "
            "as a simplifying assumption, and Supplement S2 describes its consequences."
        ),
        57: (
            "Revised Introduction: “The model assumes conditional independence among antigen-isotype "
            "responses given time since infection. Positive residual correlation can "
            "cause the current product likelihood to overstate precision. Supplement S2 "
            "provides further details.”"
        ),
        86: "See Supplement S4.1 for additional details.",
        88: "Excerpt from S4.1, Latent Infection Time:",
        102: "We added further detail to Supplement S6 (formerly Supplement S3):",
        110: (
            "Response: Thank you for pointing this out. The Lotka–Volterra system "
            "motivates the two-phase model, but serocalculator's incidence likelihood "
            "uses the antibody decay phase after pathogen concentration has reached zero. "
            "Because cross-sectional surveys primarily reflect past infections, the brief "
            "growth phase and pathogen trajectory do not enter the final decay likelihood."
        ),
        119: (
            "Response: Thank you for this thoughtful observation. We agree that random "
            "effects on kinetic parameters can produce time-varying variance around the "
            "population mean trajectory. It is important, however, to distinguish this "
            "marginal population variance from residual observation noise conditional on "
            "an individual's kinetic parameters. The upstream serodynamics model uses a "
            "constant residual variance on the log-antibody scale for each antigen-isotype, "
            "while allowing subject-specific random effects—including the waning rate. "
            "serocalculator then averages each likelihood contribution over Monte Carlo "
            "draws from the kinetic-parameter distribution rather than evaluating one "
            "mean trajectory. It therefore retains trajectory-induced variation without "
            "fitting an explicit residual-variance function of time."
        ),
        128: (
            "Following the reviewer's suggestion, we plotted posterior predictive "
            "variance against time since infection. For the fitted parameter distribution "
            "used in our example, variance was greatest during the acute phase and "
            "decreased later as the fitted power-law trajectories converged. This pattern "
            "is specific to the joint parameter distribution and should not be treated as "
            "a universal property: random waning rates can produce either divergence or "
            "convergence depending on the kinetic model, parameter scales, and correlations. "
            "The key point is that serocalculator integrates over the resulting distribution "
            "of trajectories. Its separate additive biological-noise and multiplicative "
            "measurement-noise widths are not explicit functions of time. We added this "
            "conditional-versus-marginal distinction to the methodology documentation and "
            "Supplement S4.3."
        ),
        177: (
            "Response: We agree that this is a limitation. serocalculator is best suited "
            "to settings where a time-averaged force of infection is meaningful. We now "
            "state this explicitly in the Introduction and discuss the assumption and "
            "its implications in Supplement S2. We also address this point in our "
            "response to Referee 1's Major Comment 2."
        ),
        179: (
            "Revised Introduction: “The model assumes a constant, homogeneous force of infection (a "
            "Poisson process). In settings with strong seasonality, epidemics, or secular "
            "trends, the estimate should be interpreted as a time-averaged rate. "
            "Supplement S2 describes this and other assumptions.”"
        ),
        184: (
            "Response: Thank you. Referee 1 raised the same issue. We agree that "
            "conditional independence may not hold across antigen-isotype responses. If "
            "positive residual correlation remains after conditioning on infection time, "
            "the product likelihood can count correlated biomarkers as independent "
            "information and yield confidence intervals that are too narrow. We now state "
            "this consequence in Supplement S2 and are developing methods that model the "
            "correlation directly."
        ),
        186: (
            "Revised Introduction: “Antigen-isotype responses are assumed conditionally independent "
            "given time since infection. Residual correlation may cause the current "
            "likelihood to overstate precision; Supplement S2 discusses this limitation.”"
        ),
        191: (
            "Response: Thank you. We now distinguish methodological requirements from "
            "model assumptions and summarize the likely consequences of violating each "
            "assumption in Supplement S2."
        ),
        195: (
            "Response: Thank you. We agree that the need for separately estimated, "
            "transportable seroresponse parameters limits applicability. Once parameters "
            "have been estimated for a pathogen and assay, users do not necessarily need "
            "their own longitudinal cohort, but they must justify transportability to "
            "their target population. We now state this requirement more clearly and "
            "discuss its implications in Supplement S2. We avoid listing a fixed set of "
            "available pathogens because the OSF repository is evolving. Future work will "
            "evaluate sensitivity to transportability violations."
        ),
        200: (
            "Line 135: “Further details on longitudinal seroresponse parameters and the "
            "companion package serodynamics are provided in Supplement S5.”"
        ),
        214: (
            "We revised the Introduction to summarize the published comparison with "
            "blood-culture surveillance while avoiding interpretation of clinical "
            "surveillance as a complete gold standard:"
        ),
        216: (
            "“This method has previously been used to estimate the seroincidence of "
            "enteric fever, a systemic bacterial infection caused by Salmonella enterica "
            "serovars Typhi and Paratyphi. Although the resulting seroincidence estimates "
            "were substantially higher than blood-culture-based clinical incidence, they "
            "preserved the same rank order of clinical burden across sites 12.”"
        ),
        221: (
            "The likelihood incorporates between-person heterogeneity in antibody "
            "kinetics by averaging each participant's contribution over the supplied "
            "Monte Carlo sample of kinetic-parameter sets from the longitudinal model, "
            "rather than conditioning on one mean seroresponse curve."
        ),
        223: (
            "The Hessian-based standard errors and confidence intervals are calculated "
            "from this marginalized likelihood surface. They therefore reflect "
            "cross-sectional sampling variability under a model that includes the "
            "supplied distribution of kinetic heterogeneity. However, they are conditional "
            "on the finite set of supplied kinetic draws and on fixed noise parameters; "
            "they are not a formal repeated-sampling propagation of uncertainty from "
            "estimating the longitudinal model or the noise parameters. We have clarified "
            "this distinction to avoid overstating interval coverage and identify fuller "
            "uncertainty propagation as future work."
        ),
        230: (
            "Line 230: “In contrast, serocalculator is designed to estimate "
            "population-level seroincidence by pairing quantitative antibody responses "
            "with modeled antibody decay curves and marginalizing over between-person "
            "kinetic heterogeneity. It supports multiple antigen-isotype pairs and "
            "separately models cross-reactivity or non-specific binding and assay "
            "measurement error. Collectively, these packages expand access to "
            "seroepidemiologic tools, with serocalculator filling a distinct niche in "
            "cross-sectional serosurveillance.”"
        ),
        237: (
            "Response: Thank you. We added the Poisson-process assumption to the Abstract, "
            "clarified the assumptions in the Introduction, and describe their "
            "implications in Supplement S2."
        ),
        239: (
            "Line 39: “Implementation: serocalculator is an open-source R package that "
            "uses a likelihood-based framework incorporating modeled antibody decay, "
            "between-person kinetic heterogeneity, cross-reactivity, and measurement "
            "noise to estimate seroincidence rates under a Poisson infection process "
            "from cross-sectional serological data.”"
        ),
        241: (
            "Line 93: “The model assumes a constant, homogeneous force of infection (a "
            "Poisson process), no durable post-infection immunity, transportability of "
            "the longitudinal seroresponse distribution, and conditional independence "
            "among antigen-isotype responses. Supplement S2 provides further details.”"
        ),
        259: (
            "Response: This issue was also raised by Referee 1. We agree that conditional "
            "independence may not hold for all diseases. Positive residual correlation "
            "can cause the product likelihood to overstate precision, so independence is "
            "not necessarily conservative. We now state this consequence in Supplement "
            "S2 and are developing methods that model cross-isotype correlation directly."
        ),
        261: (
            "Line 93: “Antigen-isotype responses are assumed conditionally independent "
            "given time since infection. Residual correlation can lead to confidence "
            "intervals that are too narrow; see Supplement S2.”"
        ),
        277: (
            "Response: Thank you for pointing this out. We agree that calibration requires "
            "more explanation than the manuscript permits, so we added guidance to "
            "Supplement S3."
        ),
        279: (
            "Line 115: “All data must use the same assay, units, and calibration or "
            "normalization procedure. Further requirements, formatting guidance, and "
            "example inputs are provided in Supplements S1 and S3.”"
        ),
        294: (
            "Response: We agree that these assumptions require clarification. "
            "serocalculator does not assume normally distributed biological noise; it "
            "uses Uniform(0, nu), with nu commonly set to the 95th percentile of a "
            "suitable negative-control distribution. Teunis and van Eijkeren (2020) show "
            "that estimates are sensitive to the noise width but comparatively robust to "
            "distributional shape. The 95th percentile itself is an adopted convention, "
            "not an optimized cutoff. Suitable controls may come from geographically "
            "distinct low-transmission populations or archived pre-epidemic samples when "
            "local unexposed controls are unavailable. Supplement S6 now explains this "
            "and recommends sensitivity analysis."
        ),
        296: (
            "Excerpt from Supplement S6: “Biological noise, nu, represents additive "
            "signal from cross-reactivity or non-specific binding. It is modeled as "
            "Uniform(0, nu), with nu commonly estimated as the 95th percentile in suitable "
            "negative controls. The model is sensitive to the width but comparatively "
            "robust to the distributional shape; the 95th percentile is a convention "
            "rather than an optimized cutoff. Sensitivity analysis over plausible values "
            "of nu is recommended.”"
        ),
        301: (
            "Response: The supplied measurement-noise parameter eps can vary by "
            "antigen-isotype and by declared strata such as laboratory. Within each row "
            "of the noise-parameter input, it is fixed and does not vary by observation "
            "or plate. eps is the bound on uniform relative error, not the CV itself; "
            "under the package model, eps = sqrt(3) × CV."
        ),
        306: (
            "Response: The Hessian-based confidence intervals are derived from the "
            "likelihood after averaging over the supplied kinetic-parameter draws, so the "
            "observation model includes the supplied distribution of between-person "
            "kinetic heterogeneity. The intervals remain conditional on that finite draw "
            "set and on fixed noise parameters; they do not formally propagate all "
            "uncertainty from fitting the longitudinal model or estimating noise. We are "
            "developing sensitivity-analysis tools for plausible noise-parameter ranges."
        ),
        309: (
            "Response: We agree that sample-size guidance is useful, but no single "
            "minimum applies across research questions and biomarker distributions. We "
            "recommend simulation to assess precision for an existing design or determine "
            "the sample size needed for a target effect. The sim_pop_data() function "
            "supports this workflow and is demonstrated in the Simulation Studies article. "
            "We added this guidance to Supplement S3."
        ),
        313: (
            "“Sample Size: We do not recommend one universal minimum because the required "
            "sample depends on the research question, biomarker distributions, and "
            "seroresponse heterogeneity. Users should simulate from their intended design "
            "to assess precision or determine a sample size for a target effect. The "
            "sim_pop_data() function and Simulation Studies article support this process.”"
        ),
        318: "Response: Thank you for pointing this out. We have added the interval bounds.",
        325: (
            "Response: We are glad you find the companion package promising. The packages "
            "remain standalone, but serodynamics exports kinetic-parameter draws in the "
            "format consumed by serocalculator. A shared vignette illustrating the full "
            "longitudinal-to-cross-sectional workflow is planned (GitHub Issue #542). Now "
            "that serodynamics is available on CRAN and has its own website, the manuscript "
            "and Supplement S5 clarify the connection."
        ),
    }
    revision_id = apply_paragraph_replacements(
        document, replacements, next_revision_id(document)
    )
    revision_id = replace_text_fragment(
        document.paragraphs[281], "End of Supplement S2", "End of Supplement S3",
        revision_id
    )
    revision_id = replace_text_fragment(
        document.paragraphs[311], "End of Supplement 2", "End of Supplement S3",
        revision_id
    )
    insertions, deletions = save_and_validate(document, output)
    return output, insertions, deletions


def main() -> None:
    results = [review_manuscript(), review_response_letter(), review_supplement()]
    for path, insertions, deletions in results:
        print(f"{path}\tinsertions={insertions}\tdeletions={deletions}")


if __name__ == "__main__":
    main()
